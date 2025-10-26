# app/services/file_generator.py

import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

def create_docx(notes_data, topic):
    """
    Generates a .docx file in memory from the notes and Q&A data.
    """
    try:
        document = Document()
        
        # --- Title ---
        title = document.add_heading(f"Comprehensive Notes on: {topic.title()}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph() # Add some space

        # --- Notes Section ---
        document.add_heading('Key Concepts and Notes', level=1)
        document.add_paragraph(notes_data.get('notes', 'No notes generated.'))

        # --- Q&A Section ---
        document.add_heading('Questions and Answers', level=1)
        
        qa_data = notes_data.get('questions', {})
        
        # Define the order of question marks
        mark_categories = ['1_mark', '2_marks', '4_marks', '6_marks', '8_marks', '10_marks']

        for category in mark_categories:
            questions = qa_data.get(category, [])
            if questions:
                # Add a heading for the mark category
                heading_text = category.replace('_', ' ').title()
                document.add_heading(f"{heading_text} Questions", level=2)
                
                for i, qa in enumerate(questions, 1):
                    # Add the question
                    q_para = document.add_paragraph()
                    q_para.add_run(f"Q{i}: {qa.get('question', '')}").bold = True
                    
                    # Add the answer
                    document.add_paragraph(qa.get('answer', ''))
                    document.add_paragraph() # Add space between questions

        # Save to a virtual file in memory
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
        
    except Exception as e:
        print(f"Error creating DOCX: {e}")
        return None

def create_pdf(notes_data, topic):
    """
    Generates a .pdf file in memory from the notes and Q&A data.
    """
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # --- Title ---
        pdf.set_font("Helvetica", "B", 20)
        pdf.cell(0, 10, f"Comprehensive Notes on: {topic.title()}", ln=True, align='C')
        pdf.ln(10) # Add space

        # --- Notes Section ---
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, 'Key Concepts and Notes', ln=True)
        pdf.set_font("Helvetica", "", 12)
        pdf.multi_cell(0, 5, notes_data.get('notes', 'No notes generated.'))
        pdf.ln(5)

        # --- Q&A Section ---
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, 'Questions and Answers', ln=True)
        
        qa_data = notes_data.get('questions', {})
        mark_categories = ['1_mark', '2_marks', '4_marks', '6_marks', '8_marks', '10_marks']

        for category in mark_categories:
            questions = qa_data.get(category, [])
            if questions:
                # Add a heading for the mark category
                heading_text = category.replace('_', ' ').title()
                pdf.set_font("Helvetica", "B", 14)
                pdf.cell(0, 10, f"{heading_text} Questions", ln=True)
                
                for i, qa in enumerate(questions, 1):
                    # Add the question
                    pdf.set_font("Helvetica", "B", 12)
                    pdf.multi_cell(0, 5, f"Q{i}: {qa.get('question', '')}")
                    pdf.ln(2)
                    
                    # Add the answer
                    pdf.set_font("Helvetica", "", 12)
                    pdf.multi_cell(0, 5, f"A: {qa.get('answer', '')}")
                    pdf.ln(5) # Add space between questions

        # Save to a virtual file in memory
        pdf_output = pdf.output(dest='S')
        file_stream = io.BytesIO(pdf_output.encode('latin-1')) # fpdf outputs latin-1 string
        file_stream.seek(0)
        
        return file_stream
        
    except Exception as e:
        print(f"Error creating PDF: {e}")
        return None